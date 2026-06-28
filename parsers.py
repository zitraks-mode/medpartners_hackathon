import pandas as pd
import pdfplumber
from docx import Document
from typing import List, Dict, Any
import re



# ─────────────────────────── XLSX ───────────────────────────

def parse_xlsx(file_path: str) -> List[Dict[str, Any]]:
    """Парсит XLSX/XLS, обходит все листы, ищет заголовок в первых 20 строках."""
    extracted_items = []
    try:
        xls = pd.ExcelFile(file_path)
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name, header=None)

            # Ищем строку-заголовок в первых 20 строках
            header_row_index = None
            for idx, row in df.head(20).iterrows():
                row_str = ' '.join(str(val).lower() for val in row.values)
                if ('услуга' in row_str or 'наименование' in row_str or 'название' in row_str) \
                        and ('цена' in row_str or 'стоимость' in row_str or 'тариф' in row_str):
                    header_row_index = idx
                    break

            # Перечитываем с правильным заголовком
            if header_row_index is not None:
                df = pd.read_excel(xls, sheet_name=sheet_name, header=header_row_index)
            else:
                df = pd.read_excel(xls, sheet_name=sheet_name)

            df.dropna(how='all', inplace=True)
            df.dropna(axis=1, how='all', inplace=True)
            df.columns = [str(c).lower().strip() for c in df.columns]

            service_col = None
            price_col = None
            price_resident_col = None
            # FIX: нерезидентских колонок может быть несколько (СНГ/ближнее зарубежье,
            # дальнее зарубежье и т.п.) — собираем все, а цену берём как максимум среди них.
            price_nonresident_cols = []

            for col in df.columns:
                if any(k in col for k in ('услуга', 'наименование', 'название')):
                    service_col = col
                    continue

                # FIX: реальные прайсы редко пишут буквально "резидент"/"нерезидент" —
                # обычно это "граждан РК" против "СНГ/ближнего зарубежья" / "дальнего зарубежья"
                # / "иностранных граждан". Матчим по этим формулировкам.
                is_nonresident = any(k in col for k in (
                    'нерезидент', 'снг', 'ближнего зарубежья', 'дальнего зарубежья',
                    'иностранн', 'не проживающих',
                ))
                is_resident = any(k in col for k in (
                    'резидент', 'граждан республики казахстан', 'граждан рк',
                    'постоянно проживающих',
                ))

                if is_nonresident and not is_resident:
                    price_nonresident_cols.append(col)
                elif is_resident and not is_nonresident:
                    price_resident_col = col
                elif any(k in col for k in ('цена', 'стоимость', 'тариф')):
                    if price_col is None:
                        price_col = col

            if not service_col:
                continue

            for _, row in df.iterrows():
                service_name = str(row[service_col]).strip()
                if not service_name or service_name.lower() in ('nan', 'none', ''):
                    continue

                price = _clean_price(row.get(price_col))
                price_resident = _clean_price(row.get(price_resident_col)) if price_resident_col else None

                # FIX: берём максимум среди всех найденных нерезидентских колонок
                # (например, и СНГ, и дальнее зарубежье — берём более высокую ставку)
                nonresident_candidates = [
                    p for p in (_clean_price(row.get(c)) for c in price_nonresident_cols) if p is not None
                ]
                price_nonresident = max(nonresident_candidates) if nonresident_candidates else None

                # Если отдельных колонок нет — кладём общую цену в resident
                if price_resident is None and price is not None:
                    price_resident = price

                extracted_items.append({
                    "service_name_raw": service_name,
                    "price": price_resident,
                    "price_resident": price_resident,
                    "price_nonresident": price_nonresident,
                })

    except Exception as e:
        print(f"[parse_xlsx] Ошибка {file_path}: {e}")

    return extracted_items


# ─────────────────────────── DOCX ───────────────────────────

def parse_docx(file_path: str) -> List[Dict[str, Any]]:
    extracted_items = []
    try:
        doc = Document(file_path)

        # 1. Пытаемся читать таблицы (если они есть)
        for table in doc.tables:
            if not table.rows:
                continue

            # FIX: ищем заголовочную строку и определяем колонки по смыслу
            # (название/код/резидент/нерезидент/общая цена), а не по жёсткой позиции
            # cells[0,1,2] — старая логика ломалась на таблицах другой ширины/порядка.
            service_col = None
            code_col = None
            price_col = None
            price_resident_col = None
            price_nonresident_cols = []
            header_found = False

            # Заголовок обычно среди первых строк таблицы
            for idx, row in enumerate(table.rows[:5]):
                cells_lower = [cell.text.strip().lower() for cell in row.cells]
                row_has_service = any(
                    any(k in c for k in ('услуга', 'наименование', 'название')) for c in cells_lower
                )
                row_has_price = any(
                    any(k in c for k in ('цена', 'стоимость', 'тариф')) for c in cells_lower
                )
                if row_has_service and row_has_price:
                    for i, c in enumerate(cells_lower):
                        if any(k in c for k in ('услуга', 'наименование', 'название')):
                            service_col = i
                        elif any(k in c for k in ('код',)):
                            code_col = i
                        elif any(k in c for k in (
                            'нерезидент', 'снг', 'ближнего зарубежья', 'дальнего зарубежья',
                            'иностранн', 'не проживающих',
                        )):
                            price_nonresident_cols.append(i)
                        elif any(k in c for k in (
                            'резидент', 'граждан республики казахстан', 'граждан рк',
                            'постоянно проживающих',
                        )):
                            price_resident_col = i
                        elif any(k in c for k in ('цена', 'стоимость', 'тариф')):
                            if price_col is None:
                                price_col = i
                    header_found = True
                    header_row_idx = idx
                    break

            if header_found and service_col is not None:
                for row in table.rows[header_row_idx + 1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) <= service_col:
                        continue
                    name = cells[service_col]
                    if not name or "раздел" in name.lower():
                        continue

                    code = cells[code_col] if code_col is not None and code_col < len(cells) else None

                    price_resident = (
                        _clean_price(cells[price_resident_col])
                        if price_resident_col is not None and price_resident_col < len(cells)
                        else None
                    )
                    nonresident_candidates = [
                        p for p in (
                            _clean_price(cells[c]) for c in price_nonresident_cols if c < len(cells)
                        ) if p is not None
                    ]
                    price_nonresident = max(nonresident_candidates) if nonresident_candidates else None

                    price_general = (
                        _clean_price(cells[price_col])
                        if price_col is not None and price_col < len(cells)
                        else None
                    )
                    if price_resident is None and price_general is not None:
                        price_resident = price_general

                    if price_resident is None and price_nonresident is None:
                        continue

                    extracted_items.append({
                        "code": code,
                        "name": name,
                        "service_name_raw": name,
                        "price": price_resident,
                        "price_resident": price_resident,
                        "price_nonresident": price_nonresident,
                    })
            else:
                # FIX: фоллбэк на старую позиционную логику для таблиц без
                # распознаваемого заголовка (например, только 3 колонки: код/название/цена)
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) < 3:
                        continue

                    code, name, raw_price = cells[0], cells[1], cells[2]
                    if "код" in code.lower() or not code:
                        continue
                    if "раздел" in name.lower() and not raw_price.strip():
                        continue

                    price = _clean_price(raw_price)
                    if price:
                        extracted_items.append({
                            "code": code,
                            "name": name,
                            "service_name_raw": name,
                            "price": price,
                            "price_resident": price,
                            "price_nonresident": None,
                        })

        # 2. Если таблиц нет, пытаемся читать как CSV-текст
        if not extracted_items:
            for para in doc.paragraphs:
                text = para.text.strip()
                if text.count(',') >= 2:
                    parts = text.split(',')
                    code = parts[0].strip()
                    raw_price = parts[-1].strip()
                    name = ",".join(parts[1:-1]).strip().replace('"', '')
                    
                    price = _clean_price(raw_price)
                    if price:
                        extracted_items.append({
                            "code": code,
                            "name": name,
                            "service_name_raw": name,
                            "price": price
                        })

    except Exception as e:
        print(f"Ошибка при чтении docx {file_path}: {e}")

    return extracted_items


# ─────────────────────────── PDF (цифровой текст) ───────────────────────────

def parse_pdf(file_path: str) -> List[Dict[str, Any]]:
    """Парсит текстовый PDF через pdfplumber с умным поиском колонок и фоллбэком на текст."""
    extracted_items = []
    try:
        with pdfplumber.open(file_path) as pdf:
            # Глобальные индексы колонок (сохраняются при переходе между страницами)
            service_idx = None
            price_idx = None
            price_resident_idx = None
            price_nonresident_idx = None

            for page in pdf.pages:
                tables = page.extract_tables()
                
                # ── ФОЛЛБЭК: Если таблицы нет (невидимые границы ячеек) ──
                if not tables:
                    text = page.extract_text()
                    if text:
                        raw_lines = [l.strip() for l in text.split('\n')]

                        # FIX: когда название услуги длинное, оно переносится на новую строку,
                        # а сама цена оказывается одна на следующей строке (например:
                        # "U2.1.2 Забор крови" \n "3 900"). Без склейки такая строка
                        # безвозвратно теряется, т.к. в ней нет числа в конце.
                        # Склеиваем строку без числа в конце со следующей строкой,
                        # если та состоит ТОЛЬКО из числа (т.е. это перенесённая цена).
                        merged_lines = []
                        skip_next = False
                        for i, line in enumerate(raw_lines):
                            if skip_next:
                                skip_next = False
                                continue

                            has_trailing_price = re.search(
                                r'[\s\-–]+\d[\d\s]*(?:[.,]\d+)?\s*(?:тг|тенге|kzt)?$',
                                line, re.IGNORECASE
                            )
                            is_lone_number_next = False
                            if not has_trailing_price and i + 1 < len(raw_lines):
                                next_line = raw_lines[i + 1]
                                is_lone_number_next = bool(
                                    re.fullmatch(r'\d[\d\s]*(?:[.,]\d+)?\s*(?:тг|тенге|kzt)?', next_line, re.IGNORECASE)
                                )

                            if is_lone_number_next:
                                merged_lines.append(f"{line} {raw_lines[i + 1]}")
                                skip_next = True
                            else:
                                merged_lines.append(line)

                        for line in merged_lines:
                            line = line.strip()
                            if not line:
                                continue
                            # FIX: колонтитулы и строки-разделы без цены
                            line_lower = line.lower()
                            if re.match(r'^(страниц[аы]|page)\s*\d+', line_lower):
                                continue
                            if re.match(r'^(блок|раздел|подраздел|секци[яи])\b', line_lower) \
                                    and not re.search(r'\d{4,}', line):
                                continue
                            price_match = re.search(r'[\s\-–]+(\d[\d\s]*(?:[.,]\d+)?)\s*(?:тг|тенге|kzt)?$', line, re.IGNORECASE)
                            if price_match:
                                price_str = price_match.group(1)
                                service_name = line[:price_match.start()].strip(' .-–')
                                
                                service_name = re.sub(r'^[A-Za-zА-Яа-я0-9.]+\s+', '', service_name).strip()
                                price = _clean_price(price_str)
                                
                                if price and len(service_name) > 3:
                                    extracted_items.append({
                                        "service_name_raw": service_name,
                                        "price": price,
                                        "price_resident": price,
                                        "price_nonresident": None,
                                    })
                    continue

                # ── ОБРАБОТКА ТАБЛИЦ ──
                for table in tables:
                    for row in table:
                        row = [str(c).strip().replace('\n', ' ') if c else '' for c in row]
                        
                        if not any(row):
                            continue

                        # FIX: фильтр колонтитулов и строк-разделов.
                        # Клиника 4 и подобные документы повторяют заголовок таблицы
                        # на каждой странице, а также содержат строки вида «Страница N»
                        # и «Раздел X. Название» — всё это не является позицией прайса.
                        joined_row = ' '.join(row).strip()
                        joined_lower = joined_row.lower()

                        # «Страница N» / «Page N» — колонтитул
                        if re.match(r'^(страниц[аы]|page)\s*\d+', joined_lower):
                            continue
                        # Строка-раздел без цены: «Блок А. Поликлиника», «Раздел 1. ...»
                        if re.match(r'^(блок|раздел|подраздел|секци[яи]|section)\b', joined_lower) \
                                and not any(_clean_price(c) for c in row):
                            continue
                        # Строка состоит только из одной ячейки с текстом — заголовок секции
                        non_empty = [c for c in row if c]
                        if len(non_empty) == 1 and not re.search(r'\d{3,}', non_empty[0]):
                            continue

                        # 1. Пытаемся найти или обновить заголовки колонок
                        row_lower = [c.lower() for c in row]
                        is_header = False
                        for i, h in enumerate(row_lower):
                            if any(k in h for k in ('услуга', 'наименование', 'название')):
                                service_idx = i
                                is_header = True
                            elif any(k in h for k in (
                                'нерезидент', 'снг', 'ближнего зарубежья', 'дальнего зарубежья',
                                'иностранн', 'не проживающих',
                            )):
                                price_nonresident_idx = i
                                is_header = True
                            elif any(k in h for k in (
                                'резидент', 'граждан республики', 'граждан рк',
                                'постоянно проживающих', 'страховых компани',
                                'оралман',
                            )):
                                # FIX: берём первую резидентскую колонку (страховые / граждане РК)
                                if price_resident_idx is None:
                                    price_resident_idx = i
                                is_header = True
                            elif any(k in h for k in ('цена', 'стоимость', 'тариф')):
                                price_idx = i
                                is_header = True
                            # единица измерения — не ценовая колонка
                            elif any(k in h for k in ('единиц', 'ед.', 'измерени')):
                                is_header = True  # отмечаем как заголовок, но индекс не сохраняем
                                
                        if is_header:
                            continue

                        # 2. Если заголовки на этой странице/документе не определились, применяем эвристику
                        current_service_idx = service_idx
                        current_price_idx = price_idx
                        
                        if current_service_idx is None or current_price_idx is None:
                            if len(row) >= 3:
                                current_service_idx = 1
                                current_price_idx = len(row) - 1
                            elif len(row) == 2:
                                current_service_idx = 0
                                current_price_idx = 1
                            else:
                                continue

                        # 3. Извлекаем данные
                        if current_service_idx < len(row) and current_price_idx < len(row):
                            service_name = row[current_service_idx].strip()
                            
                            if not service_name or service_name.lower() in ('none', 'nan', '') or len(service_name) < 3:
                                continue

                            # FIX: пропускаем строки где «название» — это «первичная»/«повторная»
                            # (Клиника 4: строки-подтипы без самостоятельного смысла)
                            if re.fullmatch(
                                r'(первичн\w*|повторн\w*|видеозвонок|посещение|услуга|прием)',
                                service_name.lower()
                            ):
                                continue
                                
                            price = _clean_price(row[current_price_idx])
                            price_res = _clean_price(row[price_resident_idx]) if price_resident_idx is not None and price_resident_idx < len(row) else None
                            price_nonres = _clean_price(row[price_nonresident_idx]) if price_nonresident_idx is not None and price_nonresident_idx < len(row) else None

                            if price_res is None and price is not None:
                                price_res = price

                            # FIX: Клиника 4 — строки с подтипом (первичная/повторная)
                            # хранят тип в отдельной колонке рядом с названием.
                            # Склеиваем их, если следующая непустая ячейка — подтип.
                            if price_resident_idx is not None:
                                for col_i, col_val in enumerate(row):
                                    if col_i in (current_service_idx, current_price_idx,
                                                 price_resident_idx,
                                                 price_nonresident_idx if price_nonresident_idx else -1):
                                        continue
                                    subtype = col_val.strip().lower()
                                    if re.fullmatch(
                                        r'(первичн\w*|повторн\w*|видеозвонок)',
                                        subtype
                                    ) and subtype:
                                        service_name = f"{service_name} ({col_val.strip()})"
                                        break

                            if price_res is not None:
                                extracted_items.append({
                                    "service_name_raw": service_name,
                                    "price": price_res,
                                    "price_resident": price_res,
                                    "price_nonresident": price_nonres,
                                })

    except Exception as e:
        print(f"[parse_pdf] Ошибка {file_path}: {e}")

    return extracted_items


# ─────────────────────────── PDF (скан / OCR) ───────────────────────────


def _ocr_page_columnar(img, lang: str = 'rus+eng') -> List[Dict[str, Any]]:
    """
    Поколоночный разбор одной страницы скана через pytesseract.image_to_data.

    Алгоритм:
    1. Получаем bbox каждого слова с координатами (x, y, w, h).
    2. Кластеризуем слова по вертикальным полосам (столбцам) на основе x-координат.
    3. Автоматически определяем колонки: «название» (самая широкая левая), «цена резидента»,
       «цена нерезидента/партнёра» (правые числовые колонки).
    4. Группируем слова в строки по близости y-координат и собираем текст каждой ячейки.
    5. Применяем фильтрацию по разумному диапазону цен.

    Это позволяет корректно разобрать таблицы вида:
        №  | Название услуги           | Код     | Материал   | Цена (резид.) | Цена (партнёр)
        1  | Общий анализ крови (ОАК)  | B02.110 | кровь ЭДТА | 1 880         | 1 410
    даже когда OCR читает всю строку одной строкой текста.
    """
    import pytesseract
    import pandas as pd

    data = pytesseract.image_to_data(
        img, lang=lang,
        output_type=pytesseract.Output.DATAFRAME,
        config='--psm 6',  # единый блок текста — лучше для таблиц
    )

    # Оставляем только слова с уверенностью > 30 и непустым текстом
    data = data[(data['conf'] > 30) & (data['text'].notna())]
    data['text'] = data['text'].astype(str).str.strip()
    data = data[data['text'] != '']
    if data.empty:
        return []

    page_width = img.width

    # ── 1. Кластеризация слов по столбцам ──────────────────────────────────
    # Делим ширину страницы на зоны: левая треть — название, правые зоны — цены.
    # Медицинские прайсы КЗ обычно: [№][Название][Код][Материал][ЦенаРез][ЦенаНерез]
    # Числовые колонки почти всегда в правой половине страницы.
    #
    # FIX: старое значение 0.45 позволяло колонке «Материал» (правее центра)
    # попасть в right_tokens, а значит её текст «кровь с ЭДТА» иногда
    # смешивался с числами. Поднимаем порог до 0.60 — на сканах КЗ-прайсов
    # ценовые колонки всегда правее 60% ширины страницы.
    RIGHT_ZONE_START = page_width * 0.60   # правее 60% — потенциальные ценовые колонки
    LEFT_ZONE_END    = page_width * 0.55   # левее 55%  — потенциально название

    def is_price_token(text: str) -> bool:
        """True если токен похож на цену (цифры с возможными пробелами/запятыми)."""
        cleaned = re.sub(r'[\s,.]', '', text)
        return bool(re.fullmatch(r'\d{3,}', cleaned))

    def split_merged_prices(tokens: list) -> list:
        """
        FIX: OCR иногда сливает две ценовые ячейки в одну длинную строку
        (напр. «18801410» вместо «1880» и «1410»). Если токен состоит только из цифр
        и длиннее 7 знаков (т.е. > 9 999 999 KZT) — пытаемся разбить по середине.
        Обе половины должны быть >= 3 цифр и каждая <= MAX_PRICE.
        """
        MAX_PRICE = 10_000_000
        result = []
        for t in tokens:
            digits_only = re.sub(r'[\s,.]', '', t)
            if re.fullmatch(r'\d{8,}', digits_only):
                # Пробуем разбить на две части примерно по середине
                mid = len(digits_only) // 2
                # Ищем разбивку вокруг середины, при которой обе части разумны
                split_done = False
                for offset in range(0, mid + 1):
                    for pos in (mid - offset, mid + offset):
                        if 3 <= pos <= len(digits_only) - 3:
                            left_val = float(digits_only[:pos])
                            right_val = float(digits_only[pos:])
                            if 0 < left_val <= MAX_PRICE and 0 < right_val <= MAX_PRICE:
                                result.append(digits_only[:pos])
                                result.append(digits_only[pos:])
                                split_done = True
                                break
                    if split_done:
                        break
                if not split_done:
                    result.append(t)  # не удалось разбить — оставляем, _clean_price отфильтрует
            else:
                result.append(t)
        return result

    def token_x_center(row) -> float:
        return row['left'] + row['width'] / 2

    # ── 2. Группировка слов в строки (±10px по вертикали) ──────────────────
    # Сортируем по y, затем группируем по близкому y_top
    data_sorted = data.sort_values(['top', 'left']).reset_index(drop=True)
    rows_by_y: Dict[int, list] = {}
    ROW_TOLERANCE = max(8, int(img.height * 0.012))  # ~1.2% высоты страницы

    for _, word in data_sorted.iterrows():
        y = int(word['top'])
        matched_key = None
        for key in rows_by_y:
            if abs(key - y) <= ROW_TOLERANCE:
                matched_key = key
                break
        if matched_key is None:
            matched_key = y
            rows_by_y[matched_key] = []
        rows_by_y[matched_key].append(word)

    items = []

    for y_key in sorted(rows_by_y.keys()):
        words_in_row = sorted(rows_by_y[y_key], key=lambda w: w['left'])
        full_text = ' '.join(w['text'] for w in words_in_row)

        # Пропускаем заголовки и секции
        lower = full_text.lower()
        if any(k in lower for k in (
            'наименование', 'услуга', 'название', 'тест', 'тариф',
            'гематология', 'биохимия', 'иммуно', 'паразит', 'микробиол',
            'эндоскоп', 'рентген', 'узи', 'функц', 'гистол',
        )):
            if not any(is_price_token(w['text']) for w in words_in_row):
                continue  # чистый заголовок без цен — пропускаем

        # Разделяем токены на «левые» (название/код) и «правые» (цены)
        left_tokens = []
        right_tokens = []
        for w in words_in_row:
            xc = token_x_center(w)
            if xc < RIGHT_ZONE_START:
                left_tokens.append(w['text'])
            else:
                right_tokens.append(w['text'])

        # Правые числовые токены — кандидаты на цены
        raw_price_candidates = [t for t in right_tokens if is_price_token(t)]
        # FIX: разбиваем склеенные OCR-ом токены перед извлечением цен
        price_candidates = split_merged_prices(raw_price_candidates)

        if not price_candidates:
            # Попытка найти цену в конце всей строки (простые двухколоночные прайсы)
            all_tokens = [w['text'] for w in words_in_row]
            tail_prices = [t for t in all_tokens[-3:] if is_price_token(t)]
            if not tail_prices:
                continue
            price_candidates = tail_prices
            # Название — всё кроме хвостовых цен
            n_tail = len(tail_prices)
            name_tokens = all_tokens[:-n_tail]
        else:
            name_tokens = left_tokens

        # Собираем название, очищаем от порядковых номеров и мед. кодов в начале
        service_name = ' '.join(name_tokens).strip()
        # Убираем ведущий порядковый номер: "1 Общий анализ..." → "Общий анализ..."
        service_name = re.sub(r'^\d{1,3}[\s.]+', '', service_name)
        # Убираем медицинский код в начале: "B02.110 Общий анализ" → "Общий анализ"
        service_name = re.sub(
            r'^[A-Za-zА-Яа-яBВ]\d{2}[.\-]\d{3}(?:[.\-]\d{3})?\s*', '', service_name
        ).strip()
        # Убираем мусорный хвост — код/материал после длинного названия (если прилип)
        # напр. "Общий анализ крови (ОАК без СОЭ) B02.110.002 кровь с ЭДТА" → чистим до скобки/слов
        service_name = re.sub(
            r'\s+[A-Z]\d{2}[.\-]\d{3}.*$', '', service_name
        ).strip()

        if len(service_name) < 4:
            continue

        # FIX: фильтруем колонтитулы и строки-разделы в OCR-тексте
        sn_lower = service_name.lower()
        if re.match(r'^(страниц[аы]|page)\s*\d*$', sn_lower):
            continue
        if re.match(r'^(блок|раздел|подраздел|секци[яи])\b', sn_lower) and not prices:
            continue

        # Берём первые две цены (резидент, нерезидент/партнёр)
        prices = []
        for pt in price_candidates[:2]:
            p = _clean_price(pt)
            if p and p > 0:
                prices.append(p)

        if not prices:
            continue

        price_resident = prices[0]
        price_nonresident = prices[1] if len(prices) > 1 else None

        items.append({
            "service_name_raw": service_name,
            "price": price_resident,
            "price_resident": price_resident,
            "price_nonresident": price_nonresident,
        })

    # FIX: Клиника 5 и подобные — строки-подтипы «первичный» / «повторный»
    # появляются как самостоятельные позиции без нормального названия.
    # Склеиваем их с предыдущей «именованной» позицией.
    SUBTYPE_RE = re.compile(
        r'^(первичн\w+|повторн\w+|видеозвонок|онлайн\b)',
        re.IGNORECASE,
    )
    merged_items = []
    last_named: dict | None = None
    for item in items:
        name = item["service_name_raw"]
        if SUBTYPE_RE.match(name) and last_named is not None:
            # Это подтип — создаём новую позицию с именем «Название (подтип)»
            merged_items.append({
                **item,
                "service_name_raw": f"{last_named['service_name_raw']} ({name})",
            })
        else:
            last_named = item
            merged_items.append(item)

    return merged_items


def parse_pdf_scan(file_path: str) -> List[Dict[str, Any]]:
    """
    OCR-парсинг сканированного PDF.

    Использует поколоночный разбор (image_to_data) вместо построчного (image_to_string),
    что позволяет корректно разделять название услуги и цены даже когда OCR
    склеивает несколько ячеек одной строки в единый текст.
    """
    extracted_items = []
    try:
        import pytesseract
        from pdf2image import convert_from_path

        # dpi=250: баланс между точностью распознавания и скоростью.
        # При dpi=200 Tesseract иногда путает '1' и '7' в ценах.
        images = convert_from_path(file_path, dpi=250)

        for page_num, img in enumerate(images, 1):
            try:
                page_items = _ocr_page_columnar(img, lang='rus+eng')
                extracted_items.extend(page_items)
                print(f"[parse_pdf_scan] Страница {page_num}: {len(page_items)} позиций")
            except Exception as e:
                print(f"[parse_pdf_scan] Ошибка на стр.{page_num}: {e}")
                # Фоллбэк: старый построчный метод если columnar не сработал
                try:
                    text = pytesseract.image_to_string(img, lang='rus+eng', config='--psm 6')
                    for line in text.splitlines():
                        line = line.strip()
                        if len(line) < 5:
                            continue
                        m = re.search(r'[\s\-–]+(\d[\d\s]{2,6}(?:[.,]\d+)?)\s*(?:тг|тенге|kzt)?$', line, re.IGNORECASE)
                        if m:
                            p = _clean_price(m.group(1))
                            name = line[:m.start()].strip(' .-–')
                            name = re.sub(r'^\d{1,3}[\s.]+', '', name)
                            name = re.sub(r'^[A-Za-zА-Яа-яBВ]\d{2}[.\-]\d{3}.*?\s+', '', name).strip()
                            if p and len(name) > 3:
                                extracted_items.append({
                                    "service_name_raw": name,
                                    "price": p,
                                    "price_resident": p,
                                    "price_nonresident": None,
                                })
                except Exception as e2:
                    print(f"[parse_pdf_scan] Фоллбэк тоже не сработал на стр.{page_num}: {e2}")

    except Exception as e:
        print(f"[parse_pdf_scan] Критическая ошибка {file_path}: {e}")

    return extracted_items


# ─────────────────────────── Определение типа PDF ───────────────────────────

def _is_scanned_pdf(file_path: str) -> bool:
    """Возвращает True если PDF не содержит текстового слоя (скан)."""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                # Если хотя бы на одной странице есть осмысленный текст — это цифровой PDF
                if text and len(text.strip()) > 20:
                    return False
        return True
    except Exception:
        return False


# ─────────────────────────── Утилиты ───────────────────────────

def _clean_price(value) -> float | None:
    """
    Очищает значение цены и возвращает float или None.

    FIX: добавлен верхний порог MAX_PRICE. Когда OCR склеивает две ценовые
    ячейки в одну строку (напр. «1 880 1 410» → «18801410»), результат
    после strip-пробелов выходит за любой реальный диапазон медицинских услуг.
    Если число превышает порог — возвращаем None вместо фантастической суммы.
    Порог 10 000 000 KZT (~22 000 USD) покрывает даже самые дорогие операции.
    """
    MAX_PRICE = 10_000_000  # KZT; всё выше — артефакт OCR
    if value is None:
        return None
    try:
        clean = str(value).replace(' ', '').replace('\xa0', '').replace(',', '.')
        clean = re.sub(r'[^\d.]', '', clean)
        if not clean:
            return None
        result = float(clean)
        if result <= 0 or result > MAX_PRICE:
            return None
        return result
    except (ValueError, TypeError):
        return None


# ─────────────────────────── Роутер ───────────────────────────

def parse_file(file_path: str) -> List[Dict[str, Any]]:
    ext = file_path.lower().rsplit('.', 1)[-1]
    if ext in ('xlsx', 'xls'):
        return parse_xlsx(file_path)
    elif ext == 'docx':
        return parse_docx(file_path)
    elif ext == 'pdf':
        if _is_scanned_pdf(file_path):
            print(f"[Router] Файл {file_path} определен как СКАН. Запуск OCR...")
            return parse_pdf_scan(file_path)
        else:
            print(f"[Router] Файл {file_path} определен как ЦИФРОВОЙ ТЕКСТ. Запуск pdfplumber...")
            return parse_pdf(file_path)
    return []